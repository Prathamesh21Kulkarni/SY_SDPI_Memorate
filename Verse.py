from docx import Document
def verse(input):
    input_para = input
    # input_para = '''Artificial intelligence (AI), is intelligence demonstrated by machines, unlike the natural intelligence displayed by humans and animals. Leading AI textbooks define the field as the study of "intelligent agents": any device that perceives its environment and takes actions that maximize its chance of successfully achieving its goals.[3] Colloquially, the term "artificial intelligence" is often used to describe machines (or computers) that mimic "cognitive" functions that humans associate with the human mind, such as "learning" and "problem solving". '''
    # print(len(input_para.split()))
    total_words = len(input_para.split())
    total_lines = total_words // 8
    # print(total_lines)
    list_of_words = []
    input_para = input_para.replace(",", '')
    input_para = input_para.replace('''"''', '')
    input_para = input_para.replace(".", '')
    input_para = input_para.replace("'", '')
    input_para = input_para.replace(":", "")
    input_para = input_para.replace("(", ' ')
    input_para = input_para.replace(")", "")
    input_para = input_para.replace("[", ' ')
    input_para = input_para.replace("]", '')
    stop_words = ["a", "an", "the"]
    Digits = ["0", "1", "2", "3", "4", "5", "6", "7", "8", "9"]
    for each_digit in Digits:
        if each_digit in input_para:
            input_para = input_para.replace(each_digit, " ")
    rhyme = open("Rhyme_words.csv", "r")
    suf = rhyme.readlines()
    list_of_suffix = []
    for each_word in suf:
        list_of_suffix.append(each_word.replace('\n', '').lower())
    list_of_suffix.pop(0)
    list_of_suffix = tuple(list_of_suffix)
    # list_of_suffix = ('able', 'ac', 'acity', 'ocity', 'ade', 'age', 'aholic', 'oholic', 'al', 'algia', 'an', 'ian', 'ance',
    #                   'ant', 'ar', 'ard', 'rian', 'rium', 'orium', 'ary', 'ate', 'tion', 'ative', 'cide', 'cracy',
    #                   'crat', 'cule', 'cy', 'cycle', 'dom', 'dox', 'ectomy', 'ed', 'ee', 'eer', 'emia', 'en', 'ence',
    #                   'ency', 'ent', 'er', 'ern', 'escence', 'ese', 'esque', 'ess', 'est', 'etic', 'ette', 'ful', 'fy',
    #                   'gam', 'gamy', 'gon', 'graphy', 'gonic', 'hood', 'ial', 'ian'
    #                   'ile', 'ily', 'ine', 'ing', 'ion', 'ious', 'ish', 'ism', 'ist', 'ite', 'itis', 'ity', 'ive',
    #                   'ization', 'ize', 'less', 'let', 'like', 'ling', 'loger', 'logist', 'log', 'ly', 'ment', 'ness',
    #                   'oid', 'ology', 'oma', 'onym', 'opia', 'opsy', 'or', 'ory', 'osis', 'ostomy', 'otomy', 'es', 's', 'iasis', 'iatric', 'ible', 'ic',
    #                    'ical',...)
    for words in input_para.split():
        if words not in stop_words:
            if words.endswith(list_of_suffix) == True and len(words) > 7:
                for each_suffix in list_of_suffix:
                    if words.endswith(each_suffix) == True:
                        words = words.rstrip(each_suffix)
                        list_of_words.append(words)
                        break
            else:
                list_of_words.append(words)
    # Extracting the words for Rhyme
    temp_words_for_rhyme = {}
    words_for_rhmye = {}
    for i in range(len(list_of_words)):
        if len(list_of_words[i]) >= 5:
            temp_words_for_rhyme[i] = list_of_words[i]
    temp_words_for_rhyme_keys = temp_words_for_rhyme.keys()
    temp_words_for_rhyme_keys = list(temp_words_for_rhyme_keys)
    prev_index = 0
    words_for_rhyme_keys = []
    for i in range(len(temp_words_for_rhyme_keys)):
        diff = temp_words_for_rhyme_keys[i] - prev_index
        if 4 <= diff <= 6:
            # words_for_rhyme_keys.append(prev_index)
            words_for_rhyme_keys.append(temp_words_for_rhyme_keys[i])
            prev_index = temp_words_for_rhyme_keys[i]
        elif diff == 3:
            words_for_rhyme_keys.append(temp_words_for_rhyme_keys[i])
            prev_index = temp_words_for_rhyme_keys[i]
        elif i == len(temp_words_for_rhyme_keys) - 1:
            words_for_rhyme_keys.append(temp_words_for_rhyme_keys[i])
    for index in words_for_rhyme_keys:
        words_for_rhmye[index] = list_of_words[index]
    # # Reading and getting nouns from Nouns.csv
    # f = open("Nouns1.csv")
    # nouns_file = f.readlines()
    # nouns = []
    # for each_line in nouns_file:
    #     noun = each_line
    #     noun = noun.replace("\n", '')
    #     nouns.append(noun)
    # nouns.pop(0)
    # nouns1 = nouns[:100]
    # #Searching nouns in given paragraph
    # nouns_in_para = {}
    # for words in input_para.split():
    #     if words not in stop_words:
    #         for i in range(len(nouns)):
    #             if words.lower() == nouns[i]:
    #                 noun_word = words
    #                 nouns_in_para[input_para.split().index(words)] = noun_word
    # Extracting the words for rhyme


    # print(list_of_words)
    # words_for_rhmye = {}
    # for i in range(0, len(list_of_words), 8):
    #     if i + 9 <= len(list_of_words):
    #         words_for_rhmye[i + 3] = list_of_words[i + 3]
    #         words_for_rhmye[i + 7] = list_of_words[i + 7]
    words_for_rhmye_key = list(words_for_rhmye.keys())
    words_for_rhmye_val = list(words_for_rhmye.values())
    vowels = ('a', 'e', 'i', 'o', 'u')
    final_words = []
    dict_rhyme_words = {}

    for i in range(0, len(words_for_rhmye_val), 2):
        if i + 1 <= len(words_for_rhmye_val) - 1:
            if len(words_for_rhmye_val[i]) <= 3:
                suffix_rhyme = words_for_rhmye_val[i]
                word_to_be_changed = words_for_rhmye_val[i + 1]
                temp = word_to_be_changed.endswith(list_of_suffix)
                if temp == True:
                    for each_suffix in list_of_suffix:
                        if word_to_be_changed.endswith(each_suffix) == True:
                            part_to_be_removed = each_suffix
                else:
                    part_to_be_removed = word_to_be_changed[-2:]
                final_word = word_to_be_changed.rstrip(part_to_be_removed)
                if (final_word.endswith(vowels) == True and suffix_rhyme.startswith(vowels) == True) or (
                        len(final_word) >= 2 and len(suffix_rhyme) >= 2):
                    if final_word[-1] == suffix_rhyme[0]:
                        final_word = final_word.rstrip(final_word[-1])
                final_rhyme_word = final_word + suffix_rhyme
                dict_rhyme_words[final_rhyme_word] = words_for_rhmye_key[i + 1]
            elif len(words_for_rhmye_val[i + 1]) <= 3:
                suffix_rhyme = words_for_rhmye_val[i + 1]
                word_to_be_changed = words_for_rhmye_val[i]
                temp = word_to_be_changed.endswith(list_of_suffix)
                if temp == True:
                    for each_suffix in list_of_suffix:
                        if word_to_be_changed.endswith(each_suffix) == True:
                            part_to_be_removed = each_suffix
                else:
                    part_to_be_removed = word_to_be_changed[-2:]
                final_word = word_to_be_changed.rstrip(part_to_be_removed)
                if (final_word.endswith(vowels) == True and suffix_rhyme.startswith(vowels) == True) or (
                        len(final_word) >= 2 and len(suffix_rhyme) >= 2):
                    if final_word[-1] == suffix_rhyme[0]:
                        final_word = final_word.rstrip(final_word[-1])
                final_rhyme_word = final_word + suffix_rhyme
                dict_rhyme_words[final_rhyme_word] = words_for_rhmye_key[i]
            else:
                word_not_to_be_changed = words_for_rhmye_val[i]
                if word_not_to_be_changed.endswith(list_of_suffix) == True:
                    for each_suffix in list_of_suffix:
                        if word_not_to_be_changed.endswith(each_suffix) == True:
                            suffix = each_suffix
                else:
                    suffix = word_not_to_be_changed[-2:]
                word_to_be_changed = words_for_rhmye_val[i + 1]
                if word_to_be_changed.endswith(list_of_suffix) == True:
                    for each_suffix in list_of_suffix:
                        if word_to_be_changed.endswith(each_suffix) == True:
                            part_to_be_removed = each_suffix
                else:
                    part_to_be_removed = word_to_be_changed[-2:]
                final_word = word_to_be_changed.rstrip(part_to_be_removed)
                if (final_word.endswith(vowels) == True and suffix.startswith(vowels) == True) or (
                        len(final_word) >= 2 and len(suffix) >= 2):
                    if final_word[-1] == suffix[0]:
                        final_word = final_word.rstrip(final_word[-1])
                final_rhyme_word = word_to_be_changed + suffix
                dict_rhyme_words[final_rhyme_word] = words_for_rhmye_key[i + 1]
    for key in dict_rhyme_words:
        list_of_words.insert(int(dict_rhyme_words[key]), key)
        list_of_words.pop(int(dict_rhyme_words[key]) + 1)
    each_line_lst = []
    prev_i = 0
    line = []
    for i in range(1,len(words_for_rhyme_keys),2):
        line.append(words_for_rhyme_keys[i])
    for i in line:
            each_line_lst.append(str(list_of_words[prev_i:i+1]))
            prev_i = i+1
    return(each_line_lst)
    # doc = Document()
    # doc.add_heading(heading, 0)
    # i = 0
    #
    # for each_line in each_line_lst:
    #     each_line = each_line.replace('[', '')
    #     each_line = each_line.replace(']', '')
    #     each_line = each_line.replace("'", '')
    #     each_line = each_line.replace(',', '')
    #     each_line = each_line.replace('''"''', '')
    #     doc.add_paragraph("| " + each_line + " ||")
    # doc.save(save_as + '.docx')
    # print("hi")
